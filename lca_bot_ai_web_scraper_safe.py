
import os
from openai import OpenAI
import openai
import requests
from bs4 import BeautifulSoup
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import datetime
import random

# Safely load API key
openai_api_key = os.getenv("OPENAI_API_KEY")
if not openai_api_key:
    st.warning("⚠️ OPENAI_API_KEY not set. Proceeding in fallback mode.")
client = OpenAI(api_key=openai_api_key) if openai_api_key else None

def generate_lci_data():
    return pd.DataFrame({
        'Life Cycle Stage': ['Materials', 'Manufacturing', 'Use Phase', 'End-of-Life'],
        'Energy Use (MJ)': [random.uniform(80, 120), random.uniform(50, 100), random.uniform(10, 20), random.uniform(15, 30)],
        'GHG Emissions (kg CO2-eq)': [random.uniform(5, 10), random.uniform(8, 12), random.uniform(1, 3), random.uniform(2, 4)],
        'Water Use (L)': [random.uniform(20, 40), random.uniform(10, 30), random.uniform(1, 5), random.uniform(5, 15)]
    })

def create_visuals(df):
    chart_files = []
    for column in df.columns[1:]:
        fig, ax = plt.subplots()
        ax.bar(df['Life Cycle Stage'], df[column], color='steelblue')
        ax.set_title(f'{column} by Stage')
        file = f"{column.replace(' ', '_')}.png"
        fig.savefig(file)
        chart_files.append(file)
        plt.close(fig)
    return chart_files

def scrape_product_data(product):
    query = product.replace(" ", "+") + "+environmental+impact"
    url = f"https://www.google.com/search?q={query}"
    headers = {"User-Agent": "Mozilla/5.0"}
    res = requests.get(url, headers=headers)
    soup = BeautifulSoup(res.text, "html.parser")
    paragraphs = soup.find_all("div", class_="BNeawe s3v9rd AP7Wnd")
    combined = " ".join([p.get_text() for p in paragraphs[:5]])
    return combined if combined.strip() else f"No public data found for {product}."

def generate_ai_section(prompt, product):
    if not client:
        return f"[Fallback] This section would normally be generated using AI for: {prompt} on {product}."
    try:
        response = client.chat.completions.create(
            model="gpt-4-turbo",
            messages=[
                {"role": "system", "content": "You are a sustainability analyst writing ISO-style LCA reports."},
                {"role": "user", "content": f"Write the '{prompt}' section for a life cycle assessment of a {product}."}
            ],
            temperature=0.7
        )
        return response.choices[0].message.content
    except openai.OpenAIError as e:
        return f"[Error or Quota Reached] {prompt} section fallback: AI generation failed due to: {e}"

def create_report(product, df, charts, web_data, ai_sections):
    doc = Document()
    doc.add_heading(f"LCA Report for: {product}", 0)
    doc.add_paragraph(f"Date: {datetime.date.today()}")
    doc.add_paragraph("Confidential – For Internal Use Only").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_page_break()

    doc.add_heading("Table of Contents", level=1)
    toc = ["Executive Summary", "1. Introduction", "2. Goal and Scope", "3. Functional Unit", "4. System Boundary",
           "5. Web-Sourced Product Information", "6. Inventory Analysis", "7. LCIA with Charts", "8. Interpretation",
           "9. Limitations", "10. Recommendations", "Appendix A: Glossary", "Appendix B: References"]
    for section in toc:
        doc.add_paragraph(section)
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(ai_sections["Executive Summary"])
    doc.add_page_break()

    for section in ["1. Introduction", "2. Goal and Scope", "3. Functional Unit", "4. System Boundary"]:
        doc.add_heading(section, level=1)
        doc.add_paragraph(ai_sections[section])
        doc.add_page_break()

    doc.add_heading("5. Web-Sourced Product Information", level=1)
    doc.add_paragraph(web_data)
    doc.add_page_break()

    doc.add_heading("6. Inventory Analysis", level=1)
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = col
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(round(val, 2)) if isinstance(val, (int, float)) else str(val)
    doc.add_page_break()

    doc.add_heading("7. LCIA with Charts", level=1)
    for chart in charts:
        doc.add_paragraph(f"Figure: {chart.split('.')[0].replace('_', ' ').title()}")
        doc.add_picture(chart, width=Inches(5.5))
    doc.add_page_break()

    for section in ["8. Interpretation", "9. Limitations", "10. Recommendations"]:
        doc.add_heading(section, level=1)
        doc.add_paragraph(ai_sections[section])
        doc.add_page_break()

    doc.add_heading("Appendix A: Glossary", level=1)
    doc.add_paragraph("LCA: Life Cycle Assessment\nGWP: Global Warming Potential\nMJ: Megajoules\nCO2-eq: Carbon dioxide equivalent")
    doc.add_page_break()

    doc.add_heading("Appendix B: References", level=1)
    doc.add_paragraph("1. ISO 14040/44\n2. Ecoinvent\n3. IPCC\n4. Manufacturer Reports\n5. Online product research")
    doc.add_page_break()

    filename = f"LCA_Report_AI_{product.replace(' ', '_')}.docx"
    doc.save(filename)
    return filename

st.title("🌿 ISO LCA Bot")

product = st.text_input("Enter a product name:", "Electric Toothbrush")

if st.button("Generate Report"):
    with st.spinner("Building LCI data, charts, and narrative..."):
        df = generate_lci_data()
        charts = create_visuals(df)
        web_data = scrape_product_data(product)

        ai_sections = {}
        for section in ["Executive Summary", "1. Introduction", "2. Goal and Scope",
                        "3. Functional Unit", "4. System Boundary", "8. Interpretation",
                        "9. Limitations", "10. Recommendations"]:
            ai_sections[section] = generate_ai_section(section, product)

        report_path = create_report(product, df, charts, web_data, ai_sections)

    with open(report_path, "rb") as f:
        st.download_button("📥 Download ISO LCA Report", f, file_name=report_path)
